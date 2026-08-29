#!/usr/bin/env python3
"""Otto Doc-Writer v8 — 零空白终极版。python create_docx.py in.md out.docx"""
from __future__ import annotations
import re, sys
from datetime import datetime
from pathlib import Path

if sys.platform == 'win32':
    try: sys.stdout.reconfigure(encoding='utf-8')
    except: pass

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("pip install python-docx"); sys.exit(1)

# ─── 色彩 ────────────────────────────────────────────────────────
def _rgb(h): h=h.lstrip("#"); return tuple(int(h[i:i+2],16) for i in(0,2,4))
def _hex(r,g,b): return f"{max(0,min(255,int(r))):02X}{max(0,min(255,int(g))):02X}{max(0,min(255,int(b))):02X}"
def _dark(h,a): r,g,b=_rgb(h); return _hex(r*(1-a),g*(1-a),b*(1-a))
def _blend(c1,c2,r): r1,g1,b1=_rgb(c1); r2,g2,b2=_rgb(c2); return _hex(r1*(1-r)+r2*r,g1*(1-r)+g2*r,b1*(1-r)+b2*r)

def _body(base): r,g,b=_rgb(base); lum=0.299*r+0.587*g+0.114*b; return _hex(r*0.35,g*0.35,b*0.35) if lum<80 else "2D2D2D"
def _muted(base): r,g,b=_rgb(base); lum=0.299*r+0.587*g+0.114*b; return _hex(r*0.40+90,g*0.40+90,b*0.40+90) if lum<80 else "777777"

def resolve(meta):
    base=meta.get("base","0A1628"); accent=meta.get("accent","2D7DD2"); surface=meta.get("surface","F0F4F8")
    return {"theme":meta.get("theme",""),"atmo":meta.get("atmosphere",""),
        "base":base,"accent":accent,"surface":surface,
        "body":_body(base),"muted":_muted(base),
        "light_tint":_blend("FFFFFF",accent,0.04),"callout_bar":_blend("FFFFFF",accent,0.3),
        "hr":_dark(_blend(base,accent,0.5),0.5),
        "hdr_bg":"F5F5F5","hdr_text":base,"stripe":surface,
        "h_font":meta.get("heading_font","Microsoft YaHei"),
        "b_font":meta.get("body_font","Microsoft YaHei"),
        "t_sz":int(meta.get("title_size","24")),"h1_sz":int(meta.get("h1_size","16")),
        "h2_sz":int(meta.get("h2_size","13")),"b_sz":int(meta.get("body_size","11")),
        "cover":meta.get("cover","true")!="false","toc":meta.get("toc","false")=="true",
        "indent":meta.get("indent","true")!="false","margin":float(meta.get("margin","2.54")),
        "line_sp":float(meta.get("line_spacing","1.5"))}

# ─── 解析 ────────────────────────────────────────────────────────
LAYOUT_RE=re.compile(r'<!--\s*layout:\s*(\w[\w-]*)\s*-->')

def parse(text):
    meta={}; body=text.strip()
    if body.startswith("---"):
        parts=body.split("---",2)
        if len(parts)>=3:
            for line in parts[1].strip().split("\n"):
                if ":" in line and not line[0]=="#": k,_,v=line.partition(":"); meta[k.strip()]=v.strip().strip('"').strip("'")
            body=parts[2].strip()
    lines=body.split("\n"); i=0
    secs=[]; cur={"heading":"","layout":"narrative","blocks":[]}
    def save():
        if cur["blocks"]: secs.append(cur.copy())
    tbl=[]; in_t=False
    while i<len(lines):
        line=lines[i]
        if line.strip().startswith("|") and line.strip().endswith("|"):
            tbl.append(line); in_t=True; i+=1; continue
        elif in_t:
            if tbl:
                hh, rr = _tbl(tbl)
                if hh: cur["blocks"].append({"t": "table", "h": hh, "r": rr})
            tbl = []; in_t = False; continue
        if not line.strip(): i+=1; continue
        m=re.match(r"^(##)\s+(.+)$",line)
        if m:
            save(); txt=m.group(2).strip(); lm=LAYOUT_RE.search(txt)
            layout="narrative"
            if lm: layout=lm.group(1); txt=txt[:lm.start()].strip()
            cur={"heading":txt,"layout":layout,"blocks":[]}; i+=1; continue
        m=re.match(r"^(#{3,6})\s+(.+)$",line)
        if m: cur["blocks"].append({"t":"sub","lvl":len(m.group(1)),"text":m.group(2).strip()}); i+=1; continue
        m=re.match(r"^[-*+]\s+(.+)$",line)
        if m:
            items=[]
            while i<len(lines) and re.match(r"^[-*+]\s+(.+)$",lines[i]):
                items.append(re.match(r"^[-*+]\s+(.+)$",lines[i]).group(1).strip()); i+=1
            cur["blocks"].append({"t":"bullet","items":items}); continue
        m=re.match(r"^\d+[.)]\s+(.+)$",line)
        if m:
            items=[]
            while i<len(lines) and re.match(r"^\d+[.)]\s+(.+)$",lines[i]):
                items.append(re.match(r"^\d+[.)]\s+(.+)$",lines[i]).group(1).strip()); i+=1
            cur["blocks"].append({"t":"ordered","items":items}); continue
        if line.startswith("> "):
            q=[]
            while i<len(lines) and lines[i].startswith("> "):
                q.append(lines[i][2:].strip()); i+=1
            cur["blocks"].append({"t":"quote","text":" ".join(q)}); continue
        if line.strip() in ("---","***","___"):
            cur["blocks"].append({"t":"hr"}); i+=1; continue
        p=[]
        while i<len(lines) and lines[i].strip() and not lines[i].startswith("#") and \
              not re.match(r"^[-*+]\s+",lines[i]) and not re.match(r"^\d+[.)]\s+",lines[i]) and \
              not lines[i].startswith("|") and not lines[i].startswith("> ") and \
              lines[i].strip() not in ("---","***","___"):
            p.append(lines[i]); i+=1
        cur["blocks"].append({"t":"para","text":"\n".join(p)})
    if tbl:
        hh, rr = _tbl(tbl)
        if hh: cur["blocks"].append({"t": "table", "h": hh, "r": rr})
    save()
    return meta,secs

def _tbl(raw):
    if len(raw)<2: return [],[]
    h=[c.strip() for c in raw[0].strip("|").split("|")]
    rows=[]
    for line in raw[1:]:
        if re.match(r"^[\|\-\s:]+$",line.strip()): continue
        cells=[c.strip() for c in line.strip("|").split("|")]
        if cells: rows.append(cells)
    return h,rows

# ─── 渲染 ────────────────────────────────────────────────────────
class R:
    def __init__(self,t,meta):
        self.t=t; self.m=meta; self.doc=Document(); self._hc=False
        self._setup(); self._styles()

    def _c(self,k): return RGBColor.from_string(self.t[k])
    def _setup(self):
        m=Cm(self.t["margin"])
        sec=self.doc.sections[0]
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=m; sec.bottom_margin=m; sec.left_margin=m; sec.right_margin=m

    def _styles(self):
        bf=self.t["b_font"]; bs=Pt(self.t["b_sz"]); hf=self.t["h_font"]; bc=self._c("body")
        ns=self.doc.styles["Normal"]; ns.font.name=bf; ns.font.size=bs; ns.font.color.rgb=bc
        ns.paragraph_format.space_after=Pt(2); ns.paragraph_format.line_spacing=self.t["line_sp"]
        self._ea(ns,bf)
        for lv in range(1,4):
            s=self.doc.styles[f"Heading {lv}"]
            sz={1:Pt(self.t["h1_sz"]),2:Pt(self.t["h2_sz"]),3:Pt(self.t["h2_sz"]-1)}[lv]
            s.font.name=hf; s.font.size=sz; s.font.bold=True
            s.font.color.rgb=self._c("base") if lv==1 else bc
            self._ea(s,hf)
            s.paragraph_format.space_before=Pt(14 if lv==1 else 8)
            s.paragraph_format.space_after=Pt(2); s.paragraph_format.keep_with_next=True
            if lv==1:
                pPr=s.element.get_or_add_pPr()
                pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="2" w:color="{self.t["accent"]}"/></w:pBdr>'))

    def _ea(self,style,font):
        rPr=style.element.get_or_add_rPr(); rf=rPr.find(qn("w:rFonts"))
        if rf is None: rf=parse_xml(f'<w:rFonts {nsdecls("w")} />'); rPr.insert(0,rf)
        rf.set(qn("w:eastAsia"),font)

    def _rn(self,r,*,name=None,sz=None,color=None,bold=False,italic=False):
        fn=name or self.t["b_font"]; ss=sz or Pt(self.t["b_sz"]); cc=color or self._c("body")
        r.font.name=fn; r.font.size=ss; r.font.color.rgb=cc; r.bold=bold; r.italic=italic
        rPr=r._element.get_or_add_rPr(); rf=rPr.find(qn("w:rFonts"))
        if rf is None: rf=parse_xml(f'<w:rFonts {nsdecls("w")} />'); rPr.insert(0,rf)
        rf.set(qn("w:eastAsia"),fn)

    # ── 封面（纯白底 + accent 细线装饰） ──────────────────────────
    def cover(self):
        if not self.t["cover"]: return
        self._hc=True
        title=self.m.get("title",""); sub=self.m.get("subtitle","")
        author=self.m.get("author",""); date=self.m.get("date","") or datetime.now().strftime("%Y年%m月")
        accent=self.t["accent"]

        # 顶部留白
        for _ in range(3):
            p=self.doc.add_paragraph()
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            p.paragraph_format.line_spacing=Pt(8)

        # 标题（白底黑字，不用深色块）
        pt=self.doc.add_paragraph(); pt.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_before=Pt(0); pt.paragraph_format.space_after=Pt(10)
        self._rn(pt.add_run(title),name=self.t["h_font"],sz=Pt(self.t["t_sz"]),color=self._c("base"),bold=True)

        # accent 色装饰线（双线：上细下粗）
        for w,sz in [(0.5,8),(0.2,2)]:
            pl=self.doc.add_paragraph(); pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
            pl.paragraph_format.space_before=Pt(2); pl.paragraph_format.space_after=Pt(2)
            pl.paragraph_format.line_spacing=Pt(w*4)  # 最小行高防折叠
            # 加一个极小空格，防止空段落边框在 Word 中折叠消失
            r=pl.add_run(" "); r.font.size=Pt(1); r.font.color.rgb=RGBColor.from_string("FFFFFF")
            pPr=pl._element.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{accent}"/></w:pBdr>'))

        if sub:
            ps=self.doc.add_paragraph(); ps.alignment=WD_ALIGN_PARAGRAPH.CENTER
            ps.paragraph_format.space_before=Pt(12)
            self._rn(ps.add_run(sub),sz=Pt(self.t["b_sz"]+2),color=self._c("body"))
        for mt in [x for x in [author,date,self.m.get("department")] if x]:
            pm=self.doc.add_paragraph(); pm.alignment=WD_ALIGN_PARAGRAPH.CENTER
            self._rn(pm.add_run(mt),sz=Pt(self.t["b_sz"]-1),color=self._c("muted"))

        # 封面后新 section
        new=self.doc.add_section()
        m=Cm(self.t["margin"])
        new.page_width=Cm(21); new.page_height=Cm(29.7)
        new.top_margin=m; new.bottom_margin=m; new.left_margin=m; new.right_margin=m

    def _bs(self):
        return self.doc.sections[1] if self._hc and len(self.doc.sections)>1 else self.doc.sections[0]

    # ── TOC（纯静态目录） ──────────────────────────────────────
    def toc(self, secs):
        if not self.t["toc"]: return
        p=self.doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
        self._rn(p.add_run("目  录"),name=self.t["h_font"],sz=Pt(self.t["h1_sz"]),color=self._c("base"),bold=True)
        pl=self.doc.add_paragraph()
        pl._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="3" w:space="1" w:color="{self.t["accent"]}"/></w:pBdr>'))

        # 收集所有条目：chapter → (level, text)
        entries = []
        for sec in secs:
            h = sec.get("heading","")
            if h: entries.append((1, h))
            for blk in sec.get("blocks",[]):
                if blk["t"] == "sub":
                    lvl = min(blk.get("lvl", 2), 3)
                    entries.append((lvl, blk["text"]))

        if entries:
            tab_pos = 8510
            tab_xml = f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:leader="dot" w:pos="{tab_pos}"/></w:tabs>'
            accent_c = self._c("accent")
            body_c   = self._c("body")

            for lvl, text in entries:
                pe = self.doc.add_paragraph()
                pe.paragraph_format.space_before  = Pt(3 if lvl == 1 else 0)
                pe.paragraph_format.space_after   = Pt(3 if lvl == 1 else 0)
                pe.paragraph_format.line_spacing  = Pt(15 if lvl == 1 else 13)

                if   lvl == 2: pe.paragraph_format.left_indent = Cm(0.6)
                elif lvl >= 3: pe.paragraph_format.left_indent = Cm(1.2)

                pPr = pe._element.get_or_add_pPr()
                for old_tabs in pPr.findall(qn("w:tabs")): pPr.remove(old_tabs)
                pPr.append(parse_xml(tab_xml))

                # 前导文字 + <w:tab/> → dot leader 自动填满至右边界
                if lvl == 1:
                    idx = sum(1 for e in entries[:entries.index((lvl, text))] if e[0]==1) + 1
                    label = f"{idx}.  {text}"
                    self._rn(pe.add_run(label), name=self.t["h_font"],
                             sz=Pt(self.t["b_sz"] + 1), color=accent_c, bold=True)
                elif lvl == 2:
                    self._rn(pe.add_run(text), name=self.t["h_font"],
                             sz=Pt(self.t["b_sz"]), color=body_c, bold=False)
                else:
                    self._rn(pe.add_run(text), name=self.t["h_font"],
                             sz=Pt(self.t["b_sz"] - 1), color=self._c("muted"), bold=False)
                # 制表符（触发右对齐 dot leader）
                rt = pe.add_run()
                rt._element.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))

        self.doc.add_page_break()

    # ── 章节 ──────────────────────────────────────────────────
    def chapter(self,title,layout):
        self.doc.add_paragraph(title,style="Heading 1")
        layout_labels={"narrative":"通栏正文","two-column":"双栏对比","data":"数据聚焦","highlight":"核心观点","closing":"总结收尾"}
        lb=layout_labels.get(layout,layout)
        p=self.doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
        self._rn(p.add_run(f"▸ {lb}"),sz=Pt(7.5),color=self._c("muted"),italic=True)

    # ── 正文块 ────────────────────────────────────────────────
    def sub(self,text,lvl):
        s=self.doc.add_paragraph(text,style=f"Heading {min(lvl,3)}")
        s.paragraph_format.space_before=Pt(8)

    def para(self,text):
        p=self.doc.add_paragraph()
        if self.t["indent"]: p.paragraph_format.first_line_indent=Cm(0.74)
        p.paragraph_format.space_after=Pt(2)
        self._fmt(p,text)

    def _fmt(self,p,text):
        for tok in re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)",text):
            if tok.startswith("**") and tok.endswith("**"): self._rn(p.add_run(tok[2:-2]),bold=True)
            elif tok.startswith("*") and tok.endswith("*") and not tok.startswith("**"): self._rn(p.add_run(tok[1:-1]),italic=True)
            elif tok.startswith("`") and tok.endswith("`"): self._rn(p.add_run(tok[1:-1]),name="Consolas",sz=Pt(self.t["b_sz"]-1),color=RGBColor(0xC7,0x25,0x4E))
            elif tok: self._rn(p.add_run(tok))

    def bullet(self,items):
        for item in items:
            p=self.doc.add_paragraph()
            p.paragraph_format.left_indent=Cm(1); p.paragraph_format.first_line_indent=Cm(-0.5)
            p.paragraph_format.space_after=Pt(0)
            self._rn(p.add_run("•  "+item))

    def ordered(self,items):
        for idx,item in enumerate(items,1):
            p=self.doc.add_paragraph()
            p.paragraph_format.left_indent=Cm(1); p.paragraph_format.first_line_indent=Cm(-0.5)
            p.paragraph_format.space_after=Pt(0)
            self._rn(p.add_run(f"{idx}.  {item}"))

    def quote(self,text):
        p=self.doc.add_paragraph()
        p.paragraph_format.left_indent=Cm(0.8); p.paragraph_format.space_before=Pt(6)
        p.paragraph_format.space_after=Pt(6)
        pPr=p._element.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="{self.t["callout_bar"]}"/></w:pBdr>'))
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.t["light_tint"]}" w:val="clear"/>'))
        self._rn(p.add_run(text),sz=Pt(self.t["b_sz"]),color=self._c("body"))

    def table(self,hdrs,rows):
        if not hdrs: return
        cols=len(hdrs); tbl=self.doc.add_table(rows=1+len(rows),cols=cols)
        tbl.alignment=1; tbl.autofit=True
        tblPr=tbl._tbl.tblPr
        if tblPr is None: tblPr=parse_xml(f'<w:tblPr {nsdecls("w")} />')
        for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
        tblPr.append(parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/><w:left w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/><w:bottom w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/><w:right w:val="single" w:sz="2" w:space="0" w:color="DDDDDD"/><w:insideH w:val="single" w:sz="2" w:space="0" w:color="EEEEEE"/><w:insideV w:val="single" w:sz="2" w:space="0" w:color="EEEEEE"/></w:tblBorders>'))
        for j,h in enumerate(hdrs):
            c=tbl.rows[0].cells[j]; c.text=""; cp=c.paragraphs[0]; cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before=Pt(2); cp.paragraph_format.space_after=Pt(2)
            c._element.get_or_add_tcPr().append(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="30"/><w:bottom w:w="30"/><w:left w:w="60"/><w:right w:w="60"/></w:tcMar>'))
            c._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.t["hdr_bg"]}" w:val="clear"/>'))
            self._rn(cp.add_run(h),name=self.t["h_font"],sz=Pt(self.t["b_sz"]-1),color=RGBColor.from_string(self.t["hdr_text"]),bold=True)
        for i,row in enumerate(rows):
            for j,val in enumerate(row):
                if j>=cols: continue
                c=tbl.rows[i+1].cells[j]; c.text=""; cp=c.paragraphs[0]
                cp.paragraph_format.space_before=Pt(1); cp.paragraph_format.space_after=Pt(1)
                c._element.get_or_add_tcPr().append(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="20"/><w:bottom w:w="20"/><w:left w:w="60"/><w:right w:w="60"/></w:tcMar>'))
                if i%2==1: c._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.t["stripe"]}" w:val="clear"/>'))
                self._rn(cp.add_run(val),sz=Pt(self.t["b_sz"]-1))

    # ── 签名/页眉页脚 ────────────────────────────────────────
    def sig(self):
        s=self.m.get("signature_unit") or self.m.get("author") or ""
        d=self.m.get("signature_date") or self.m.get("date") or ""
        if not s and not d: return
        for line in[s,d]:
            if line: p=self.doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; self._rn(p.add_run(line))

    def hf(self):
        title=self.m.get("title",""); sec=self._bs()
        hdr=sec.header; hdr.is_linked_to_previous=False
        hp=hdr.paragraphs[0]; hp.paragraph_format.space_after=Pt(0)
        self._rn(hp.add_run(title or ""),sz=Pt(7.5),color=self._c("muted"))
        hp._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="2" w:space="1" w:color="{self.t["accent"]}"/></w:pBdr>'))
        ftr=sec.footer; ftr.is_linked_to_previous=False
        fp=ftr.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        fp._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="2" w:space="1" w:color="E0E0E0"/></w:pBdr>'))
        self._rn(fp.add_run("— "),sz=Pt(7.5),color=self._c("muted"))
        for tag in["begin",None,"end"]:
            rr=fp.add_run()
            rr._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="{tag}"/>' if tag else f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        self._rn(fp.add_run(" —"),sz=Pt(7.5),color=self._c("muted"))
        if self._hc: cs=self.doc.sections[0]; cs.different_first_page_header_footer=True; cs.header.paragraphs[0].clear(); cs.footer.paragraphs[0].clear()

    # ── 主流程 ────────────────────────────────────────────────
    def build(self,secs):
        self.doc.core_properties.title=self.m.get("title",""); self.doc.core_properties.author=self.m.get("author","")
        self.cover()
        if not self._hc:
            title=self.m.get("title","")
            if title:
                p=self.doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
                self._rn(p.add_run(title),name=self.t["h_font"],sz=Pt(self.t["t_sz"]),color=self._c("base"),bold=True)
                pl=self.doc.add_paragraph(); pl.alignment=WD_ALIGN_PARAGRAPH.CENTER
                pl._element.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="3" w:space="1" w:color="{self.t["accent"]}"/></w:pBdr>'))
        self.toc(secs)
        for sec in secs:
            if sec.get("heading"): self.chapter(sec["heading"],sec.get("layout","narrative"))
            for blk in sec.get("blocks",[]):
                t=blk["t"]
                if t=="sub": self.sub(blk["text"],blk["lvl"])
                elif t=="para": self.para(blk["text"])
                elif t=="bullet": self.bullet(blk["items"])
                elif t=="ordered": self.ordered(blk["items"])
                elif t=="quote": self.quote(blk["text"])
                elif t=="table": self.table(blk["h"],blk["r"])
                elif t=="hr": pass
            if sec.get("layout")=="closing": self.sig()
        self.sig(); self.hf()
    def save(self,p): self.doc.save(p)

def main():
    import argparse
    p=argparse.ArgumentParser(description="Otto Doc-Writer v8")
    p.add_argument("input"); p.add_argument("output")
    a=p.parse_args()
    ip=Path(a.input)
    if not ip.exists(): print(f"not found: {a.input}"); sys.exit(1)
    meta,secs=parse(ip.read_text(encoding="utf-8"))
    t=resolve(meta)
    if "title" not in meta: meta["title"]=ip.stem
    g=R(t,meta); g.build(secs)
    op=Path(a.output); op.parent.mkdir(parents=True,exist_ok=True)
    g.save(str(op))
    print(f"OK {op.name} {op.stat().st_size/1024:.0f}KB {len(secs)}sections")

if __name__=="__main__": main()
